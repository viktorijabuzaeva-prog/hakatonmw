"""
Transcript Parser Module
Extracts text content from .docx interview transcripts
"""
import os
import re
from docx import Document
from typing import Dict, List, Optional
from datetime import datetime


# Список известных банков для извлечения из текста
KNOWN_BANKS = [
    'Сбербанк', 'Сбер', 'СберБанк',
    'Тинькофф', 'Тинькоф', 'Tinkoff',
    'Альфа-Банк', 'Альфа Банк', 'Альфабанк', 'Альфа',
    'ВТБ', 'VTB',
    'Райффайзен', 'Райффайзенбанк', 'Raiffeisen',
    'Газпромбанк', 'Газпром банк',
    'Россельхозбанк', 'РСХБ',
    'Открытие', 'Банк Открытие',
    'Промсвязьбанк', 'ПСБ',
    'Совкомбанк',
    'Росбанк',
    'Почта Банк', 'Почтабанк',
    'Ренессанс', 'Ренессанс Кредит',
    'Хоум Кредит', 'Home Credit',
    'МТС Банк', 'МТС-Банк',
    'Уралсиб',
    'АК Барс', 'Ак Барс Банк',
    'Банк Санкт-Петербург',
    'Юникредит', 'UniCredit',
    'Ситибанк', 'Citibank',
    'HSBC',
    'Точка', 'Точка Банк',
    'Модульбанк', 'Модуль Банк',
    'Озон Банк', 'Ozon Bank',
    'Яндекс Банк',
]

# Нормализация названий банков к единому виду
BANK_NORMALIZATION = {
    'Сбер': 'Сбербанк',
    'СберБанк': 'Сбербанк',
    'Тинькоф': 'Тинькофф',
    'Tinkoff': 'Тинькофф',
    'Альфа Банк': 'Альфа-Банк',
    'Альфабанк': 'Альфа-Банк',
    'Альфа': 'Альфа-Банк',
    'VTB': 'ВТБ',
    'Райффайзенбанк': 'Райффайзен',
    'Raiffeisen': 'Райффайзен',
    'Газпром банк': 'Газпромбанк',
    'РСХБ': 'Россельхозбанк',
    'Банк Открытие': 'Открытие',
    'Почтабанк': 'Почта Банк',
    'Home Credit': 'Хоум Кредит',
    'МТС-Банк': 'МТС Банк',
    'Ак Барс Банк': 'АК Барс',
    'UniCredit': 'Юникредит',
    'Citibank': 'Ситибанк',
    'Точка Банк': 'Точка',
    'Модуль Банк': 'Модульбанк',
    'Ozon Bank': 'Озон Банк',
}


class TranscriptParser:
    """Parse .docx transcript files and extract structured content"""
    
    def __init__(self, transcripts_dir: str = "Transcripts"):
        self.transcripts_dir = transcripts_dir
        self._banks_cache = {}  # Кэш извлечённых банков
    
    def extract_banks(self, text: str) -> List[str]:
        """
        Извлечь упомянутые банки из текста
        
        Args:
            text: Текст для анализа
            
        Returns:
            Список уникальных банков (нормализованные названия)
        """
        found_banks = set()
        text_lower = text.lower()
        
        for bank in KNOWN_BANKS:
            # Ищем банк в тексте (без учёта регистра)
            if bank.lower() in text_lower:
                # Нормализуем название
                normalized = BANK_NORMALIZATION.get(bank, bank)
                found_banks.add(normalized)
        
        return sorted(list(found_banks))
    
    def parse_docx(self, file_path: str) -> Dict[str, any]:
        """
        Extract text from a .docx file
        
        Args:
            file_path: Path to the .docx file
            
        Returns:
            Dictionary containing transcript metadata and content
        """
        try:
            doc = Document(file_path)
            
            # Extract all text content
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():  # Skip empty paragraphs
                    paragraphs.append(para.text.strip())
            
            full_text = '\n\n'.join(paragraphs)
            
            # Extract filename without extension as respondent name
            filename = os.path.basename(file_path)
            respondent_name = os.path.splitext(filename)[0]
            
            # Get file stats
            file_stats = os.stat(file_path)
            
            return {
                'respondent_name': respondent_name,
                'file_path': file_path,
                'file_name': filename,
                'content': full_text,
                'paragraphs': paragraphs,
                'paragraph_count': len(paragraphs),
                'word_count': len(full_text.split()),
                'char_count': len(full_text),
                'file_size': file_stats.st_size,
                'modified_date': datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
                'parsed_date': datetime.now().isoformat()
            }
            
        except Exception as e:
            return {
                'error': str(e),
                'file_path': file_path,
                'success': False
            }
    
    def list_transcripts(self, extract_banks: bool = True) -> List[Dict[str, str]]:
        """
        List all .docx files in the transcripts directory
        
        Args:
            extract_banks: Если True, извлекать упомянутые банки из каждого файла
        
        Returns:
            List of dictionaries with transcript metadata
        """
        transcripts = []
        
        if not os.path.exists(self.transcripts_dir):
            return transcripts
        
        for filename in os.listdir(self.transcripts_dir):
            if filename.endswith('.docx') and not filename.startswith('~'):  # Skip temp files
                file_path = os.path.join(self.transcripts_dir, filename)
                file_stats = os.stat(file_path)
                name = os.path.splitext(filename)[0]
                
                transcript_data = {
                    'name': name,
                    'filename': filename,
                    'path': file_path,
                    'size': file_stats.st_size,
                    'modified': datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
                    'banks': []
                }
                
                # Извлекаем банки (с кэшированием)
                if extract_banks:
                    if name in self._banks_cache:
                        transcript_data['banks'] = self._banks_cache[name]
                    else:
                        try:
                            doc = Document(file_path)
                            text = '\n'.join([p.text for p in doc.paragraphs])
                            banks = self.extract_banks(text)
                            self._banks_cache[name] = banks
                            transcript_data['banks'] = banks
                        except Exception as e:
                            print(f"Error extracting banks from {filename}: {e}")
                            transcript_data['banks'] = []
                
                transcripts.append(transcript_data)
        
        # Sort by modified date (newest first)
        transcripts.sort(key=lambda x: x['modified'], reverse=True)
        
        return transcripts
    
    def get_transcript_by_name(self, name: str) -> Optional[Dict[str, any]]:
        """
        Get a specific transcript by respondent name
        
        Args:
            name: Respondent name (without .docx extension)
            
        Returns:
            Parsed transcript dictionary or None if not found
        """
        # Try exact match first
        file_path = os.path.join(self.transcripts_dir, f"{name}.docx")
        
        if os.path.exists(file_path):
            return self.parse_docx(file_path)
        
        # Try with spaces instead of dashes
        name_with_spaces = name.replace('-', ' ')
        file_path = os.path.join(self.transcripts_dir, f"{name_with_spaces}.docx")
        
        if os.path.exists(file_path):
            return self.parse_docx(file_path)
        
        # Try with dashes instead of spaces
        name_with_dashes = name.replace(' ', '-')
        file_path = os.path.join(self.transcripts_dir, f"{name_with_dashes}.docx")
        
        if os.path.exists(file_path):
            return self.parse_docx(file_path)
        
        # Try case-insensitive search
        name_lower = name.lower().replace('-', ' ')
        for filename in os.listdir(self.transcripts_dir):
            if filename.endswith('.docx'):
                file_name_lower = os.path.splitext(filename)[0].lower()
                if file_name_lower == name_lower or file_name_lower.replace(' ', '-') == name_lower.replace(' ', '-'):
                    file_path = os.path.join(self.transcripts_dir, filename)
                    return self.parse_docx(file_path)
        
        return None
    
    def parse_all_transcripts(self) -> List[Dict[str, any]]:
        """
        Parse all transcripts in the directory
        
        Returns:
            List of parsed transcript dictionaries
        """
        transcripts = self.list_transcripts()
        parsed_transcripts = []
        
        for transcript_meta in transcripts:
            parsed = self.parse_docx(transcript_meta['path'])
            if 'error' not in parsed:
                parsed_transcripts.append(parsed)
        
        return parsed_transcripts
    
    def save_transcript(self, file_data, filename: str) -> Dict[str, any]:
        """
        Save an uploaded transcript file
        
        Args:
            file_data: File binary data or file object
            filename: Desired filename (should end with .docx)
            
        Returns:
            Dictionary with save status and file info
        """
        try:
            if not filename.endswith('.docx'):
                filename += '.docx'
            
            file_path = os.path.join(self.transcripts_dir, filename)
            
            # Create directory if it doesn't exist
            os.makedirs(self.transcripts_dir, exist_ok=True)
            
            # Save file
            if hasattr(file_data, 'save'):
                # Flask FileStorage object
                file_data.save(file_path)
            else:
                # Binary data
                with open(file_path, 'wb') as f:
                    f.write(file_data)
            
            return {
                'success': True,
                'file_path': file_path,
                'filename': filename,
                'message': 'Transcript saved successfully'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'message': 'Failed to save transcript'
            }


def test_parser():
    """Test function to verify parser works"""
    parser = TranscriptParser()
    
    print("=== Testing Transcript Parser ===\n")
    
    # List all transcripts
    transcripts = parser.list_transcripts()
    print(f"Found {len(transcripts)} transcripts:\n")
    
    if transcripts:
        # Parse first transcript as example
        first = transcripts[0]
        print(f"Parsing: {first['name']}")
        
        parsed = parser.parse_docx(first['path'])
        
        if 'error' not in parsed:
            print(f"\nRespondent: {parsed['respondent_name']}")
            print(f"Paragraphs: {parsed['paragraph_count']}")
            print(f"Words: {parsed['word_count']}")
            print(f"Characters: {parsed['char_count']}")
            print(f"\nFirst 200 characters:")
            print(parsed['content'][:200] + "...")
        else:
            print(f"Error: {parsed['error']}")


if __name__ == "__main__":
    test_parser()
